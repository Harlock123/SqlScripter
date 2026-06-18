#!/usr/bin/env python3
"""Generate USAGEDOCUMENT_ForSQLSCRIPTER.DOCX from the rendered screenshots.

Run via the project venv:
    .docvenv/bin/python tools/gen_usage_doc.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOTS = os.path.join(ROOT, "screenshots")
OUT = os.path.join(ROOT, "USAGEDOCUMENT_ForSQLSCRIPTER.DOCX")

IMG_WIDTH = Inches(6.3)
ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def shot(path):
    return os.path.join(SHOTS, path)


doc = Document()

# Base style.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# ---------------------------------------------------------------- title page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SqlScripter")
run.font.size = Pt(34)
run.font.bold = True
run.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Usage Guide")
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("A cross-platform MSSQL “drop & recreate” scripting tool\n"
                "Windows · macOS · Linux  —  Intel/AMD (x64) and ARM64")
r.font.size = Pt(12)
r.italic = True

doc.add_paragraph()


def heading(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def sub_heading(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def figure(image, cap):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(shot(image), width=IMG_WIDTH)
    caption(cap)


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


# ---------------------------------------------------------------- overview
heading("1. Overview")
doc.add_paragraph(
    "SqlScripter connects to a Microsoft SQL Server database, enumerates its "
    "tables, views and stored procedures, and lets you select any combination "
    "of them to generate a single T-SQL script that drops (if present) and "
    "recreates each object. The generated script appears in a built-in, "
    "SQL-syntax-highlighted editor where it can be reviewed, copied or saved."
)
sub_heading("Key features")
bullets([
    "Connect to any reachable MSSQL instance with a standard connection string.",
    "Browse Tables, Views and Stored Procedures in a categorized tree.",
    "Multi-select objects with checkboxes; select whole categories at once.",
    "Filter the object list by a “contains” or “does not contain” pattern.",
    "Toggle existence checks & DROP statements on or off.",
    "Optionally script table data as INSERT statements, with a per-table row cap.",
    "Syntax-highlighted T-SQL output with line numbers.",
    "Copy the script to the clipboard, save it to a .sql file, or clear the output.",
    "Runs natively on Windows, macOS and Linux (x64 and ARM64).",
])

# ---------------------------------------------------------------- launch
heading("2. Launching the Application")
doc.add_paragraph(
    "On start-up the window is divided into three areas: the connection bar "
    "across the top, the object tree and scripting options on the left, and the "
    "generated-SQL editor on the right. A status bar along the bottom reports "
    "progress and results."
)
figure("01-launch.png", "Figure 1 — SqlScripter immediately after launch.")

# ---------------------------------------------------------------- connecting
heading("3. Connecting to a Database")
doc.add_paragraph(
    "Enter a connection string in the box at the top of the window and click "
    "Connect. A typical SQL-authentication connection string looks like:"
)
mono = doc.add_paragraph()
r = mono.add_run(
    "Server=localhost,1433;Database=MyDb;User Id=sa;Password=Pass123;"
    "TrustServerCertificate=True;Encrypt=True;"
)
r.font.name = "Consolas"
r.font.size = Pt(10)
doc.add_paragraph(
    "When the connection succeeds, the tree is populated and the status bar "
    "reports how many tables, views and stored procedures were found. "
    "Integrated/Windows and Azure AD authentication are also supported — the "
    "tool simply passes the connection string through to the SQL client."
)

# ---------------------------------------------------------------- browsing
heading("4. Browsing Database Objects")
doc.add_paragraph(
    "Discovered objects are grouped under three categories — Tables, Views and "
    "Stored Procedures — each showing a count. Expand a category to see its "
    "members, listed as schema-qualified names (for example, dbo.Customer)."
)
figure("02-connected.png",
       "Figure 2 — The object tree populated after a successful connection.")

# ---------------------------------------------------------------- selecting
heading("5. Selecting Objects to Script")
doc.add_paragraph(
    "Tick the checkbox next to any object to include it. Ticking a category "
    "checkbox selects all of its children at once. Use the All and None buttons "
    "to select or clear everything quickly."
)
doc.add_paragraph(
    "Selections are remembered independently of the filter: you can tick some "
    "objects, change the filter, tick more, and the Generate Script command will "
    "include everything you ticked — even rows currently hidden by the filter. "
    "The None button clears the entire selection."
)

# ---------------------------------------------------------------- filtering
heading("6. Filtering the Object List")
doc.add_paragraph(
    "Above the tree, a filter lets you narrow what is shown. Choose Contains or "
    "Does not contain from the drop-down and type a pattern; matching is "
    "case-insensitive and applied to each object’s schema.name. In the example "
    "below, the pattern “usp_” with Contains reduces the list to just the "
    "stored procedures."
)
figure("03-filter.png",
       "Figure 3 — Filtering with “Contains usp_” narrows the tree to stored procedures.")

# ---------------------------------------------------------------- options
heading("7. Scripting Options")
doc.add_paragraph("Three options on the left panel control what the generated script contains:")
bullets([
    "Script existence checks & DROP if present — emits an IF OBJECT_ID(...) "
    "guard that drops each object before it is recreated. Turn this off to "
    "produce CREATE-only scripts.",
    "Script INSERT records for tables — also reproduces table data as INSERT "
    "statements (identity columns are handled with SET IDENTITY_INSERT).",
    "Max rows per table (0 = all) — caps how many rows are scripted per table. "
    "This field is enabled only when INSERT scripting is turned on.",
])
figure("04-options.png",
       "Figure 4 — Drops and INSERT scripting enabled with a 500-row cap, three tables selected.")

# ---------------------------------------------------------------- generating
heading("8. Generating and Reviewing the Script")
doc.add_paragraph(
    "Click Generate Script to build the T-SQL for every selected object. The "
    "result is shown in the editor on the right, syntax-highlighted as MSSQL "
    "with line numbers. For each object the script contains an optional "
    "existence check and DROP followed by a full recreation:"
)
bullets([
    "Tables — CREATE TABLE with current data types, identity, computed columns "
    "and defaults, then PRIMARY KEY / UNIQUE / FOREIGN KEY / CHECK constraints "
    "and secondary indexes (plus optional INSERT data).",
    "Views — the stored CREATE VIEW definition.",
    "Stored procedures — the stored CREATE PROCEDURE definition.",
])
figure("05-generated.png",
       "Figure 5 — A generated drop & recreate script in the syntax-highlighted editor.")

# ---------------------------------------------------------------- output actions
heading("9. Copying, Saving and Clearing Output")
doc.add_paragraph(
    "The toolbar above the editor provides three actions for the generated script:"
)
bullets([
    "Copy — copies the entire script to the system clipboard.",
    "Save… — opens a file dialog to write the script to a .sql file.",
    "Clear — empties the output pane.",
])

# ---------------------------------------------------------------- platforms
heading("10. Supported Platforms & Building")
doc.add_paragraph(
    "SqlScripter is built with .NET 9 and Avalonia and runs on Windows, macOS "
    "and Linux on both Intel/AMD (x64) and ARM64. Self-contained, single-file "
    "executables can be produced for every target — no .NET install is required "
    "on the destination machine."
)
bullets([
    "macOS / Linux: run ./publish-all.sh (bash, no PowerShell required).",
    "Windows: run ./publish-all.ps1 (PowerShell).",
    "Runtime identifiers: win-x64, win-arm64, osx-x64, osx-arm64, linux-x64, linux-arm64.",
])
doc.add_paragraph(
    "Note: because it uses Microsoft.Data.SqlClient, the application requires ICU "
    "globalization data; minimal Linux containers may need the libicu package "
    "installed."
)

# ---------------------------------------------------------------- footer + page numbers
GREY = RGBColor(0x60, 0x60, 0x60)


def _add_field(run, instr):
    """Append a Word field (e.g. PAGE, NUMPAGES) to an existing run."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr_text)
    run._r.append(end)


def _top_border(paragraph):
    """Draw a thin rule above a paragraph (used to separate the footer)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), "BFBFBF")
    borders.append(top)
    p_pr.append(borders)


def _grey(run):
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    return run


section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False

fp = footer.paragraphs[0]
fp.text = ""
_top_border(fp)

# Tab stops: a right-aligned stop at the usable page width pins "Page X of Y"
# to the right margin while the title stays on the left.
usable = section.page_width - section.left_margin - section.right_margin
fp.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)

_grey(fp.add_run("SqlScripter — Usage Guide"))
page = _grey(fp.add_run("\tPage "))
_add_field(page, "PAGE")
_grey(fp.add_run(" of "))
_add_field(_grey(fp.add_run("")), "NUMPAGES")

doc.save(OUT)
print("Wrote", OUT)
